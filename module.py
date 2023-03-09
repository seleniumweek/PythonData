import docx
import PyPDF2 as pyp
import difflib
import win32com.client
from bs4 import BeautifulSoup
import pandas as pd
import numpy as np
from xlsxwriter.utility import xl_rowcol_to_cell

def getDocxTextToString(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

def readandSaveDocxDifferences(filename1,filename2):
    doc1 = docx.Document(filename1)
    doc2 = docx.Document(filename2)
    if len(doc1.paragraphs) != len(doc2.paragraphs):
        print("The number of paragraphs in the documents is different.")
    else:
        value=""
        for i in range(len(doc1.paragraphs)):
            if doc1.paragraphs[i].text != doc2.paragraphs[i].text:
                value=value+"Line number: {}\nFile1:{}\nFile2:{}\n\n".format(i,doc1.paragraphs[i].text,doc2.paragraphs[i].text)
        return value        

def readPdfFileAndConvertString(filename1):
            listvalues = []
            pdfFileObj = open(filename1, 'rb')
            pdfReader = pyp.PdfReader(pdfFileObj)
            pages= len(pdfReader.pages)
            for page in range(pages):
                eachpagepdftext =  pdfReader.pages[page].extract_text()
                listvalues.append(eachpagepdftext.splitlines())
            pdfFileObj.close()
            return listvalues

def readCSVText(filename1,filename2):
    file1 = open(filename1, "r")
    file2 = open(filename2, "r")
    i = 0
    value =""
    for l1 in file1:
        i += 1
        for l2 in file2:
            if l1 != l2:
                cur =""
                if l1=='\n':
                    cur = 'Line No: {}\nFile1: Blank Line \nFile2:{}\n\n'.format(i,l2.replace('\n',''))
                elif(l2 == '\n'):      
                    cur = 'Line No: {}\nFile1:{}\nFile2: Blank Line\n\n'.format(i,l1.replace('\n',''))
                else:
                    cur = 'Line No: {}\nFile1:{}\nFile2:{}\n\n'.format(i,l1.replace('\n',''),l2.replace('\n',''))      
                    value = value + cur
                    break
    return value

def readHTMLDocument(filename1, filename2):

    '''with open(filename1, "r") as file1:
        html1 = file1.read()
    with open(filename2, "r") as file2:
        html2 = file2.read()
    d = difflib.Differ()
    diff = d.compare(html1.splitlines(), html2.splitlines())
    return "\n".join(diff)'''

    with open(filename1) as file1, open(filename2) as file2:
        '''for file1Line, file2Line in zip(file1, file2):
            if file1Line != file2Line:
                print(file1Line.strip('\n'))
                print(file2Line.strip('\n'))'''
 
    soup = BeautifulSoup(file1)
    print(soup.get_text())

def ConvertDocToHtml(inputfile, outfilename):
    doc = win32com.client.GetObject(inputfile)
    doc.SaveAs (FileName=outfilename, FileFormat=8)
    doc.Close ()            

def readPDFDocument(filename1, filename2):

    with open(filename1, "rb") as file1:
        pdf1 = file1.read()
    with open(filename2, "rb") as file2:
        pdf2 = file2.read()
    pdf1_string = pdf1.decode("utf-8")
    pdf2_string = pdf2.decode("utf-8")
    diff = difflib.unified_diff(pdf1_string.splitlines(), pdf2_string.splitlines())
    value=""
    for line in diff:
           value = value+line

    return value     

def ReadExcelData():
    template = pd.read_excel("C:\\Users\\admin\\Desktop\\Book1.xlsx",na_values=np.nan,header=None)
    testSheet = pd.read_excel("C:\\Users\\admin\\Desktop\\Book2.xlsx",na_values=np.nan,header=None)

    rt,ct = template.shape
    rtest,ctest = testSheet.shape

    df = pd.DataFrame(columns=['Cell_Location','File1_Value','File2_Value'])

    for rowNo in range(max(rt,rtest)):
        for colNo in range(max(ct,ctest)):
            # Fetching the template value at a cell
            try:
                template_val = template.iloc[rowNo,colNo]
            except:
                template_val = np.nan

            # Fetching the testsheet value at a cell
            try:
                testSheet_val = testSheet.iloc[rowNo,colNo]
            except:
                testSheet_val = np.nan

            # Comparing the values
            if (str(template_val)!=str(testSheet_val)):
                cell = xl_rowcol_to_cell(rowNo, colNo)
                dfTemp = pd.DataFrame([[cell,template_val,testSheet_val]],
                                    columns=['Cell_Location','File1_Value','File2_Value'])
                df = df.append(dfTemp)
                #print(df)
    return df            